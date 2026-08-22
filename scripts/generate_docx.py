#!/usr/bin/env python3
"""
从 Markdown 草稿生成格式化 DOCX 论文
支持国赛(CUMCM)和美赛(MCM/ICM)模板

用法:
    python generate_docx.py <draft_md> --template cumcm|mcm \
        --output <output.docx> [--figures-dir <dir>]
"""

import argparse
import re
import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import json


# Windows 控制台默认代码页可能不是 UTF-8。论文正文不能依赖控制台区域设置。
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")


# ============================================================
# 格式配置
# ============================================================

CUMCM_CONFIG = {
    "paper_size": ("A4", Cm(21.0), Cm(29.7)),
    "margins": {"top": Cm(2.5), "bottom": Cm(2.5), "left": Cm(3.0), "right": Cm(2.5)},
    "line_spacing": 1.25,
    "fonts": {
        "title": {"name": "黑体", "size": Pt(16), "bold": True},
        "heading1": {"name": "黑体", "size": Pt(16), "bold": True},
        "heading2": {"name": "黑体", "size": Pt(14), "bold": True},
        "heading3": {"name": "黑体", "size": Pt(12), "bold": True},
        "body": {"name": "宋体", "size": Pt(12), "bold": False},
        "fig_caption": {"name": "宋体", "size": Pt(10.5), "bold": False},
        "table_caption": {"name": "宋体", "size": Pt(10.5), "bold": False},
        "reference": {"name": "宋体", "size": Pt(9), "bold": False},
        "code": {"name": "Courier New", "size": Pt(9), "bold": False},
    },
    "first_indent": Cm(0.74),  # 约 2 字符
    "fig_label_format": "图{num} {caption}",
    "table_label_format": "表{num} {caption}",
    "ref_format": "GB/T 7714",
    "abstract_keyword_label": "关键词：",
}

MCM_CONFIG = {
    "paper_size": ("Letter", Inches(8.5), Inches(11.0)),
    "margins": {"top": Inches(1.0), "bottom": Inches(1.0), "left": Inches(1.0), "right": Inches(1.0)},
    "line_spacing": 1.15,
    "fonts": {
        "title": {"name": "Times New Roman", "size": Pt(16), "bold": True},
        "heading1": {"name": "Times New Roman", "size": Pt(13), "bold": True},
        "heading2": {"name": "Times New Roman", "size": Pt(12), "bold": True},
        "heading3": {"name": "Times New Roman", "size": Pt(11), "bold": True},
        "body": {"name": "Times New Roman", "size": Pt(12), "bold": False},
        "fig_caption": {"name": "Times New Roman", "size": Pt(10), "bold": False},
        "table_caption": {"name": "Times New Roman", "size": Pt(10), "bold": False},
        "reference": {"name": "Times New Roman", "size": Pt(10), "bold": False},
        "code": {"name": "Courier New", "size": Pt(9), "bold": False},
    },
    "first_indent": Inches(0.3),
    "fig_label_format": "Figure {num}: {caption}",
    "table_label_format": "Table {num}: {caption}",
    "ref_format": "APA",
    "abstract_keyword_label": "Keywords: ",
}


def get_config(template: str) -> dict:
    if template.lower() == "mcm":
        return MCM_CONFIG
    return CUMCM_CONFIG


# ============================================================
# 文档设置
# ============================================================

def setup_document(doc: Document, config: dict):
    """设置页面、边距、默认样式"""
    # 页边距
    for section in doc.sections:
        m = config["margins"]
        section.top_margin = m["top"]
        section.bottom_margin = m["bottom"]
        section.left_margin = m["left"]
        section.right_margin = m["right"]

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    body_font = config["fonts"]["body"]
    font.name = body_font["name"]
    font.size = body_font["size"]
    font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = config["line_spacing"]

    # 设置中文字体回退
    rPr = style.element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{body_font["name"]}"/>')
    rPr.insert(0, rFonts)

    # 创建标题样式
    for level, key in [(1, "heading1"), (2, "heading2"), (3, "heading3")]:
        heading_style = doc.styles[f'Heading {level}']
        h_font = config["fonts"][key]
        heading_style.font.name = h_font["name"]
        heading_style.font.size = h_font["size"]
        heading_style.font.bold = h_font["bold"]
        heading_style.font.color.rgb = RGBColor(0, 0, 0)
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)


def add_paragraph(doc: Document, text: str, style_name: str = 'Normal',
                  bold: bool = False, alignment=None, font_name=None,
                  font_size=None, first_indent=None):
    """添加段落并设置格式"""
    para = doc.add_paragraph(style=style_name)
    run = para.add_run(text)
    run.font.color.rgb = RGBColor(0, 0, 0)

    # 仅设置 run.font.name 只覆盖西文字体；Word 在 Windows 上可能因此
    # 用错误的东亚字体替换中文。显式写入三类字体槽，沿用 v1.3 的兼容策略。
    selected_font = font_name or "宋体"
    _set_run_font_slots(run, selected_font)

    if bold:
        run.bold = True
    if font_name:
        run.font.name = font_name
    if font_size:
        run.font.size = font_size
    if alignment is not None:
        para.alignment = alignment

    if first_indent:
        para.paragraph_format.first_line_indent = first_indent

    return para


def _set_run_font_slots(run, east_asia_font: str):
    """显式设置 Word 的西文与东亚字体槽，避免依赖系统区域设置。"""
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    latin_font = "Times New Roman" if east_asia_font in {"宋体", "黑体"} else east_asia_font
    r_fonts.set(qn("w:ascii"), latin_font)
    r_fonts.set(qn("w:hAnsi"), latin_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


# ============================================================
# Markdown 解析
# ============================================================

def parse_markdown_to_sections(md_text: str) -> list[dict]:
    """
    将 Markdown 解析为结构化节列表。
    每节：{level, title, content_lines, images, tables}
    """
    lines = md_text.split('\n')
    sections = []
    current_section = None
    in_code_block = False

    for line in lines:
        stripped = line.strip()

        # 代码块
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            if current_section:
                current_section['content_lines'].append(line)
            continue

        # 标题行
        if not in_code_block and stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            title = stripped.lstrip('#').strip()

            if current_section:
                sections.append(current_section)

            current_section = {
                "level": level,
                "title": title,
                "content_lines": [],
                "images": [],
                "tables": [],
            }
            continue

        if current_section is None:
            current_section = {
                "level": 0,
                "title": "正文",
                "content_lines": [],
                "images": [],
                "tables": [],
            }

        # 图片
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', stripped)
        if img_match:
            current_section['images'].append({
                "alt": img_match.group(1),
                "path": img_match.group(2),
            })
            current_section['content_lines'].append(f"[图片: {img_match.group(1)}]")
            continue

        # 表格（记录但保留原始）
        if stripped.startswith('|'):
            current_section['tables'].append(stripped)

        current_section['content_lines'].append(line)

    if current_section:
        sections.append(current_section)

    return sections


# ============================================================
# 内容渲染
# ============================================================

def _split_content_blocks(lines: list[str]) -> list[tuple[str, object]]:
    """把节内容行切分为 (kind, payload) 序列。

    - kind 为 "para" 时，payload 是普通行文本，由调用方按空行切段落；
    - kind 为 "code" 时，payload 是代码块内的行列表（保留空行、不含围栏）。
    未闭合的代码围栏也会被保留，避免内容静默丢失。
    """
    blocks: list[tuple[str, object]] = []
    buf: list[str] = []
    code_buf: list[str] = []
    in_code = False

    for line in lines:
        stripped = line.strip()
        if stripped.startswith('```'):
            if in_code:
                blocks.append(("code", code_buf))
                code_buf = []
                in_code = False
            else:
                if buf:
                    blocks.append(("para", "\n".join(buf)))
                    buf = []
                in_code = True  # 围栏起始行（含语言标记）不渲染
            continue
        if in_code:
            code_buf.append(line)
        else:
            buf.append(line)

    if in_code:
        blocks.append(("code", code_buf))
    elif buf:
        blocks.append(("para", "\n".join(buf)))
    return blocks


def render_section(doc: Document, section: dict, config: dict, figures_dir: str):
    """将一节内容渲染到 DOCX"""
    level = section['level']
    title = section['title']

    if level >= 1 and level <= 3:
        # 使用标题样式
        add_paragraph(doc, title, style_name=f'Heading {min(level, 3)}',
                      bold=True,
                      font_name=config["fonts"][f"heading{min(level, 3)}"]["name"],
                      font_size=config["fonts"][f"heading{min(level, 3)}"]["size"])
    elif title and title != "正文":
        add_paragraph(doc, title, bold=True)

    # 渲染内容：按行扫描，代码块保留为等宽字体段落，其余按段落渲染
    for kind, payload in _split_content_blocks(section['content_lines']):
        if kind == "code":
            # 代码块：等宽字体、无首行缩进，逐行保留（含空行）
            para = doc.add_paragraph()
            run = para.add_run()
            code_lines = payload
            for i, code_line in enumerate(code_lines):
                if i > 0:
                    run.add_break()
                run.add_text(code_line)
            run.font.name = config["fonts"]["code"]["name"]
            run.font.size = config["fonts"]["code"]["size"]
            _set_run_font_slots(run, config["fonts"]["code"]["name"])
            continue

        # 普通段落：按空行切分
        for para_text in payload.split('\n\n'):
            para_text = para_text.strip()
            if not para_text:
                continue

            # 跳过纯图片标记行
            if para_text.startswith('[图片:') and len(para_text) < 100:
                continue

            # 内联 LaTeX 公式保留 $...$ 原始文本，Word 中可用公式编辑器进一步处理
            add_paragraph(doc, para_text, first_indent=config.get("first_indent"))

    # 插入图片
    for img in section.get('images', []):
        img_path = img['path']
        # 尝试解析图片路径
        if not os.path.isabs(img_path):
            img_path = os.path.join(figures_dir, os.path.basename(img_path))

        if os.path.exists(img_path):
            try:
                doc.add_picture(img_path, width=Inches(5.5))
                # 图片标题
                caption_text = img['alt'] if img['alt'] else os.path.basename(img_path)
                fig_para = add_paragraph(
                    doc, caption_text,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    font_name=config["fonts"]["fig_caption"]["name"],
                    font_size=config["fonts"]["fig_caption"]["size"],
                )
                doc.add_paragraph()  # 空行
            except Exception as e:
                add_paragraph(doc, f"[图片插入失败: {img_path} - {e}]",
                              font_name=config["fonts"]["fig_caption"]["name"],
                              font_size=config["fonts"]["fig_caption"]["size"])
        else:
            add_paragraph(doc, f"[图片未找到: {img_path}]",
                          font_name=config["fonts"]["fig_caption"]["name"],
                          font_size=config["fonts"]["fig_caption"]["size"])


# ============================================================
# 摘要页
# ============================================================

def add_summary_page(doc: Document, sections: list[dict], config: dict):
    """添加摘要页"""
    # 查找摘要段
    abstract_text = ""
    keywords = ""

    for section in sections:
        if '摘要' in section['title'] or 'Abstract' in section['title']:
            abstract_text = '\n'.join(section['content_lines']).strip()

            # 提取关键词
            kw_match = re.search(
                r'(?:关键词|Keywords)[：:]\s*(.*)',
                abstract_text
            )
            if kw_match:
                keywords = kw_match.group(1).strip()
                abstract_text = abstract_text.replace(kw_match.group(0), '').strip()

            # 在当前文档开头插入摘要
            break

    if abstract_text:
        # 摘要标题
        if config == CUMCM_CONFIG:
            add_paragraph(doc, "摘要", alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          bold=True,
                          font_name=config["fonts"]["heading1"]["name"],
                          font_size=config["fonts"]["heading1"]["size"])
        else:
            add_paragraph(doc, "Summary", alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          bold=True,
                          font_name=config["fonts"]["heading1"]["name"],
                          font_size=config["fonts"]["heading1"]["size"])

        # 摘要正文
        for para_text in abstract_text.split('\n\n'):
            para_text = para_text.strip()
            if para_text and not para_text.startswith('关键词') and not para_text.startswith('Keywords'):
                add_paragraph(doc, para_text, first_indent=config.get("first_indent"))

        # 关键词
        if keywords:
            label = config.get("abstract_keyword_label", "关键词：")
            kw_para = doc.add_paragraph()
            kw_run_label = kw_para.add_run(label)
            kw_run_label.bold = True
            kw_run_label.font.name = config["fonts"]["body"]["name"]
            kw_run_label.font.size = config["fonts"]["body"]["size"]
            kw_run_label.font.color.rgb = RGBColor(0, 0, 0)
            _set_run_font_slots(kw_run_label, config["fonts"]["body"]["name"])
            kw_run_content = kw_para.add_run(keywords)
            kw_run_content.font.name = config["fonts"]["body"]["name"]
            kw_run_content.font.size = config["fonts"]["body"]["size"]
            kw_run_content.font.color.rgb = RGBColor(0, 0, 0)
            _set_run_font_slots(kw_run_content, config["fonts"]["body"]["name"])

        # 分页
        doc.add_page_break()


# ============================================================
# 主函数
# ============================================================

def generate_docx(md_path: str, template: str, output_path: str, figures_dir: str):
    """主生成函数"""
    config = get_config(template)

    # 读取 Markdown
    source_path = Path(md_path)
    try:
        md_text = source_path.read_text(encoding="utf-8-sig")
    except UnicodeDecodeError:
        # 兼容 Windows 编辑器导出的 GBK/GB18030 草稿；输出 DOCX 仍由
        # python-docx 统一写为合法 UTF-8 XML。
        md_text = source_path.read_text(encoding="gb18030")

    # 解析结构
    sections = parse_markdown_to_sections(md_text)

    # 创建文档
    doc = Document()
    setup_document(doc, config)

    # 摘要页
    if template.lower() != "mcm":
        # 国赛摘要页在正文前
        add_summary_page(doc, sections, config)

    # 渲染各节（跳过摘要相关节，已单独处理）
    for section in sections:
        title = section.get('title', '')
        if template.lower() != "mcm" and ('摘要' in title or 'Abstract' in title):
            # 国赛摘要已单独处理
            # 但保留美赛摘要（在正文流中）
            continue
        render_section(doc, section, config, figures_dir)

    # 保存
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    print(f"论文已生成: {output_path}")
    print(f"模板: {template}")
    print(f"共 {len(sections)} 个章节")
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Markdown 转格式化 DOCX 论文")
    parser.add_argument("draft_md", help="Markdown 草稿文件路径")
    parser.add_argument("--template", required=True, choices=["cumcm", "mcm"],
                        help="论文模板: cumcm(国赛) 或 mcm(美赛)")
    parser.add_argument("--output", required=True, help="输出 DOCX 文件路径")
    parser.add_argument("--figures-dir", default=".", help="图片文件所在目录")
    args = parser.parse_args()

    if not Path(args.draft_md).exists():
        print(f"错误：草稿文件不存在 {args.draft_md}")
        sys.exit(1)

    generate_docx(args.draft_md, args.template, args.output, args.figures_dir)


if __name__ == "__main__":
    main()
