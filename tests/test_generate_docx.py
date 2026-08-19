#!/usr/bin/env python3
"""generate_docx.py Markdown→DOCX 渲染测试。

验证代码块被完整保留（等宽字体段落）、围栏标记不泄漏、
普通段落与标题正常渲染。依赖 python-docx。
"""

import importlib.util
import unittest
import zipfile
from docx import Document
from lxml import etree
from pathlib import Path

SCRIPTS = Path(__file__).resolve().parents[1] / "scripts"


def load_script(name):
    spec = importlib.util.spec_from_file_location(name.stem, name)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


class GenerateDocxTestCase(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.g = load_script(SCRIPTS / "generate_docx.py")

    def render_md(self, md_text: str) -> list[str]:
        sections = self.g.parse_markdown_to_sections(md_text)
        doc = Document()
        self.g.setup_document(doc, self.g.CUMCM_CONFIG)
        for section in sections:
            self.g.render_section(doc, section, self.g.CUMCM_CONFIG, ".")
        return [p.text for p in doc.paragraphs]

    def test_code_block_preserved_as_paragraph(self):
        md = (
            "# 标题\n\n"
            "正文段落。\n\n"
            "```python\n"
            'print("hello")\n'
            "x = 1\n"
            "```\n\n"
            "结尾段落。\n"
        )
        texts = self.render_md(md)
        self.assertTrue(any('print("hello")' in t for t in texts), texts)
        self.assertTrue(any("x = 1" in t for t in texts), texts)
        self.assertTrue(any("正文段落" in t for t in texts), texts)
        self.assertTrue(any("结尾段落" in t for t in texts), texts)
        # 围栏标记本身不应泄漏到正文
        self.assertFalse(any("```" in t for t in texts), texts)

    def test_unclosed_code_block_still_preserved(self):
        md = "# 标题\n\n```python\nprint(\"lost-line\")\n"
        texts = self.render_md(md)
        self.assertTrue(any('print("lost-line")' in t for t in texts), texts)

    def test_heading_and_abstract_flow(self):
        md = "# 摘要\n\n这是摘要内容。\n\n# 问题一\n\n内容。\n"
        texts = self.render_md(md)
        self.assertTrue(any(t.strip() == "摘要" for t in texts), texts)
        self.assertTrue(any(t.strip() == "问题一" for t in texts), texts)

    def test_chinese_text_and_east_asia_font_are_preserved(self):
        md = "# 中文标题\n\n温度为 25 ℃，误差为 ±0.1；结论：模型有效。\n"
        sections = self.g.parse_markdown_to_sections(md)
        doc = Document()
        self.g.setup_document(doc, self.g.CUMCM_CONFIG)
        for section in sections:
            self.g.render_section(doc, section, self.g.CUMCM_CONFIG, ".")

        with self.subTest("text"):
            text = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("温度为 25 ℃", text)
            self.assertIn("模型有效", text)

        with self.subTest("font mapping"):
            paragraph = next(p for p in doc.paragraphs if "模型有效" in p.text)
            r_fonts = paragraph.runs[0]._element.rPr.rFonts
            self.assertEqual(r_fonts.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"), "宋体")

    def test_gb18030_markdown_is_decoded_before_docx_generation(self):
        with self.subTest("decode"):
            import tempfile
            with tempfile.TemporaryDirectory() as tmp:
                source = Path(tmp) / "draft.md"
                output = Path(tmp) / "paper.docx"
                source.write_bytes("# 摘要\n\n关键词：优化\n\n中文正文。\n".encode("gb18030"))
                self.g.generate_docx(str(source), "mcm", str(output), tmp)
                self.assertTrue(output.is_file())
                with zipfile.ZipFile(output) as archive:
                    xml = etree.fromstring(archive.read("word/document.xml"))
                rendered = "".join(xml.xpath("//w:t/text()", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}))
                self.assertIn("中文正文", rendered)


if __name__ == "__main__":
    unittest.main()
